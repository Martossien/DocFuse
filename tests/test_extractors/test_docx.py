"""Tests de l'extracteur DOCX.

CdC §8.3 — Body, tableaux, headers/footers, footnotes, endnotes.
"""

from __future__ import annotations

import io
from pathlib import Path

import pytest

from docfuse.core.ocr.tesseract import TesseractEngine
from docfuse.extractors.docx import DocxExtractor
from docfuse.models.file_status import FileStatus

_OCR_AVAILABLE = TesseractEngine().is_available()


def _docx_with_image(tmp_path: Path, name: str, image_bytes: bytes) -> Path:
    from docx import Document

    f = tmp_path / name
    doc = Document()
    doc.add_paragraph("Texte avant image.")
    doc.add_picture(io.BytesIO(image_bytes))
    doc.save(str(f))
    return f


def _red_square_png() -> bytes:
    from PIL import Image

    buf = io.BytesIO()
    Image.new("RGB", (50, 50), "red").save(buf, format="PNG")
    return buf.getvalue()


def _inject_textbox(docx_path: Path, part_name: str, before_closing_tag: str, text: str) -> None:
    """Injecte un `w:txbxContent` minimal (VML legacy `w:pict`, comme
    produit par Word) dans une partie ZIP d'un .docx déjà sauvegardé —
    python-docx ne permet pas de créer des zones de texte par API."""
    import zipfile

    with zipfile.ZipFile(str(docx_path)) as zf:
        xml = zf.read(part_name).decode("utf-8")

    textbox_xml = (
        "<w:p><w:r><w:pict><v:shape><v:textbox><w:txbxContent>"
        f"<w:p><w:r><w:t>{text}</w:t></w:r></w:p>"
        "</w:txbxContent></v:textbox></v:shape></w:pict></w:r></w:p>"
    )
    new_xml = xml.replace(before_closing_tag, textbox_xml + before_closing_tag)
    assert new_xml != xml, f"balise de fermeture {before_closing_tag!r} introuvable"

    with zipfile.ZipFile(str(docx_path)) as zin:
        items = zin.infolist()
        contents = {item.filename: zin.read(item.filename) for item in items}
    contents[part_name] = new_xml.encode("utf-8")
    with zipfile.ZipFile(str(docx_path), "w", zipfile.ZIP_DEFLATED) as zout:
        for item in items:
            zout.writestr(item, contents[item.filename])


def _rewrite_zip(docx_path: Path, remplacements: dict[str, bytes]) -> None:
    """Réécrit un .docx en remplaçant/ajoutant les entrées données."""
    import zipfile

    with zipfile.ZipFile(str(docx_path)) as zin:
        items = zin.infolist()
        contents = {item.filename: zin.read(item.filename) for item in items}
    noms = [item.filename for item in items]
    contents.update(remplacements)
    for nom in remplacements:
        if nom not in noms:
            noms.append(nom)
    with zipfile.ZipFile(str(docx_path), "w", zipfile.ZIP_DEFLATED) as zout:
        for nom in noms:
            zout.writestr(nom, contents[nom])


def _inject_part(
    docx_path: Path, part_name: str, xml: str, content_type: str, reltype: str
) -> None:
    """Ajoute une partie XML au paquet (entrée ZIP + override de type de
    contenu + relation depuis `document.xml`) — python-docx ne charge que les
    parties atteignables par une relation, et n'a d'API ni pour les notes de
    bas de page ni pour les notes de fin."""
    import zipfile

    with zipfile.ZipFile(str(docx_path)) as zf:
        types = zf.read("[Content_Types].xml").decode("utf-8")
        rels = zf.read("word/_rels/document.xml.rels").decode("utf-8")

    types = types.replace(
        "</Types>",
        f'<Override PartName="/{part_name}" ContentType="{content_type}"/></Types>',
    )
    rid = "rIdInjecte1"
    rels = rels.replace(
        "</Relationships>",
        f'<Relationship Id="{rid}" Type="{reltype}" '
        f'Target="{part_name.split("/")[-1]}"/></Relationships>',
    )
    _rewrite_zip(
        docx_path,
        {
            part_name: xml.encode("utf-8"),
            "[Content_Types].xml": types.encode("utf-8"),
            "word/_rels/document.xml.rels": rels.encode("utf-8"),
        },
    )


_W_NS = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
_OOXML_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_OOXML_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml"


def _inject_sdt_rows(docx_path: Path, avant_marqueur: str, lignes: list[tuple[str, str]]) -> None:
    """Insère des `w:tr` enveloppés dans un `w:sdt` (`CT_SdtContentRow`, le
    contrôle « section répétitive » de Word) juste avant la ligne contenant
    `avant_marqueur` — python-docx ne sait pas produire cet encodage."""
    import zipfile

    with zipfile.ZipFile(str(docx_path)) as zf:
        xml = zf.read("word/document.xml").decode("utf-8")

    bloc = ""
    for i, (gauche, droite) in enumerate(lignes):
        cellules = "".join(
            f'<w:tc><w:tcPr><w:tcW w:w="2000" w:type="dxa"/></w:tcPr>'
            f"<w:p><w:r><w:t>{texte}</w:t></w:r></w:p></w:tc>"
            for texte in (gauche, droite)
        )
        bloc += (
            f'<w:sdt><w:sdtPr><w:id w:val="{100 + i}"/></w:sdtPr>'
            f"<w:sdtContent><w:tr>{cellules}</w:tr></w:sdtContent></w:sdt>"
        )

    debut_tr = xml.rindex("<w:tr", 0, xml.rindex(avant_marqueur))
    _rewrite_zip(
        docx_path, {"word/document.xml": (xml[:debut_tr] + bloc + xml[debut_tr:]).encode("utf-8")}
    )


class TestDocxExtractor:
    def test_extract_basic_text(self, tmp_path: Path) -> None:
        from docx import Document

        f = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph(
            "Ceci est un paragraphe de test avec suffisamment de texte "
            "pour depasser le seuil de quatre-vingt caracteres."
        )
        doc.save(str(f))

        result = DocxExtractor.extract(f, "test.docx")
        assert result.status is FileStatus.READY
        assert "paragraphe de test" in result.text
        assert result.file_type == "docx"

    def test_extract_table(self, tmp_path: Path) -> None:
        from docx import Document

        f = tmp_path / "table.docx"
        doc = Document()
        doc.add_paragraph("Texte avant le tableau " * 10)
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Nom"
        table.cell(0, 1).text = "Valeur"
        table.cell(1, 0).text = "Alpha"
        table.cell(1, 1).text = "Beta"
        doc.save(str(f))

        result = DocxExtractor.extract(f, "table.docx")
        assert result.status is FileStatus.READY
        assert "Nom" in result.text
        assert "Alpha" in result.text

    def test_accepts(self) -> None:
        assert DocxExtractor.accepts(Path("test.docx")) is True
        assert DocxExtractor.accepts(Path("test.pdf")) is False

    def test_safe_extract_no_crash(self, tmp_path: Path) -> None:
        f = tmp_path / "nonexistent.docx"
        result = DocxExtractor.safe_extract(f, "nonexistent.docx")
        assert result.status is FileStatus.ERROR

    def test_password_protected_gives_clear_error(self, tmp_path: Path) -> None:
        """D-089 : un .docx protégé par mot de passe à l'ouverture (conteneur
        OLE2, plus un ZIP) donnait un `PackageNotFoundError` bas niveau,
        jamais un message disant à l'utilisateur que le fichier est protégé."""
        f = tmp_path / "protected.docx"
        f.write_bytes(bytes((0xD0, 0xCF, 0x11, 0xE0, 0xA1, 0xB1, 0x1A, 0xE1)) + b"\x00" * 500)

        result = DocxExtractor.extract(f, "protected.docx")
        assert result.status is FileStatus.ERROR
        assert result.error_message is not None
        assert "mot de passe" in result.error_message.lower()

    def test_count_media_images_empty(self, tmp_path: Path) -> None:
        from docx import Document

        f = tmp_path / "no_images.docx"
        Document().save(str(f))
        result = DocxExtractor.extract(f, "no_images.docx")
        assert result.image_count == 0

    def test_fixture_file(self) -> None:
        fixture = Path(__file__).resolve().parent.parent / "fixtures" / "sample.docx"
        if fixture.exists():
            result = DocxExtractor.extract(fixture, "sample.docx")
            assert result.status is FileStatus.READY
            assert "Titre Test" in result.text or "Titre" in result.text

    def test_tracked_changes_insertion_is_extracted(self, tmp_path: Path) -> None:
        """D-069 : le texte inséré en suivi des modifications (w:ins) ne doit
        pas disparaître — `Paragraph.text` de python-docx ne regarde que les
        runs enfants directs de w:p, pas ceux imbriqués sous w:ins."""
        from docx import Document
        from docx.oxml.ns import qn
        from lxml import etree

        f = tmp_path / "tracked_changes.docx"
        doc = Document()
        doc.add_paragraph("Texte normal avant.")
        p = doc.add_paragraph()
        ins = etree.SubElement(p._p, qn("w:ins"))
        ins.set(qn("w:id"), "1")
        ins.set(qn("w:author"), "Test")
        r = etree.SubElement(ins, qn("w:r"))
        t = etree.SubElement(r, qn("w:t"))
        t.text = "TEXTE_INSERE_SUIVI_MODIFS"
        doc.save(str(f))

        result = DocxExtractor.extract(f, "tracked_changes.docx")
        assert result.status is FileStatus.READY
        assert "TEXTE_INSERE_SUIVI_MODIFS" in result.text

    def test_content_control_block_is_extracted(self, tmp_path: Path) -> None:
        """D-069 : un paragraphe entier enveloppé dans un contrôle de contenu
        Word (w:sdt, niveau bloc — omniprésent dans les modèles RH/juridique)
        ne doit pas disparaître : ni "}p" ni "}tbl" ne matchent "}sdt"."""
        from docx import Document
        from docx.oxml.ns import qn
        from lxml import etree

        f = tmp_path / "content_control.docx"
        doc = Document()
        doc.add_paragraph("Texte normal avant.")
        sdt = etree.SubElement(doc.element.body, qn("w:sdt"))
        sdt_content = etree.SubElement(sdt, qn("w:sdtContent"))
        inner_p = etree.SubElement(sdt_content, qn("w:p"))
        inner_r = etree.SubElement(inner_p, qn("w:r"))
        inner_t = etree.SubElement(inner_r, qn("w:t"))
        inner_t.text = "TEXTE_CONTROLE_DE_CONTENU_BLOC"
        # Le sectPr final doit rester le dernier enfant du body (contrat OOXML)
        sect_pr = doc.element.body.find(qn("w:sectPr"))
        if sect_pr is not None:
            doc.element.body.append(sect_pr)
        doc.save(str(f))

        result = DocxExtractor.extract(f, "content_control.docx")
        assert result.status is FileStatus.READY
        assert "TEXTE_CONTROLE_DE_CONTENU_BLOC" in result.text

    def test_textbox_in_body_is_extracted(self, tmp_path: Path) -> None:
        """D-082 : `find_all("w:txbxcontent")` (minuscules) ne matchait
        jamais `<w:txbxContent>` (casse réelle produite par Word, sensible
        à la casse dans le parseur XML BeautifulSoup) — l'extraction des
        zones de texte ne trouvait donc jamais rien, sur aucun fichier."""
        from docx import Document

        f = tmp_path / "textbox_body.docx"
        doc = Document()
        doc.add_paragraph("Corps normal.")
        doc.save(str(f))

        _inject_textbox(f, "word/document.xml", "</w:body>", "TEXTE_BOITE_CORPS")

        result = DocxExtractor.extract(f, "textbox_body.docx")
        assert result.status is FileStatus.READY
        assert "TEXTE_BOITE_CORPS" in result.text

    def test_textbox_in_header_is_extracted(self, tmp_path: Path) -> None:
        """D-082 : une zone de texte dans un en-tête/pied de page (logo +
        bloc adresse en papier à en-tête) était invisible — seul
        document.xml était lu, jamais word/header*.xml."""
        from docx import Document

        f = tmp_path / "textbox_header.docx"
        doc = Document()
        doc.add_paragraph("Corps normal.")
        section = doc.sections[0]
        section.header.is_linked_to_previous = False
        section.header.paragraphs[0].text = "En-tete texte normal"
        doc.save(str(f))

        _inject_textbox(f, "word/header1.xml", "</w:hdr>", "TEXTE_BOITE_EN_TETE")

        result = DocxExtractor.extract(f, "textbox_header.docx")
        assert result.status is FileStatus.READY
        assert "TEXTE_BOITE_EN_TETE" in result.text

    def test_nested_table_in_cell_is_extracted(self, tmp_path: Path) -> None:
        """D-083 : `_Cell.paragraphs` ne liste que les paragraphes directs
        d'une cellule, jamais un tableau imbriqué (fréquent dans les
        gabarits de rapports/formulaires complexes)."""
        from docx import Document

        f = tmp_path / "nested_table.docx"
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        cell.text = "TEXTE_CELLULE_EXTERNE"
        inner = cell.add_table(rows=1, cols=1)
        inner.cell(0, 0).text = "TEXTE_TABLEAU_IMBRIQUE"
        doc.save(str(f))

        result = DocxExtractor.extract(f, "nested_table.docx")
        assert result.status is FileStatus.READY
        assert "TEXTE_CELLULE_EXTERNE" in result.text
        assert "TEXTE_TABLEAU_IMBRIQUE" in result.text

    def test_table_in_header_is_extracted(self, tmp_path: Path) -> None:
        """D-107 : `_Header.paragraphs` ne rend que les `w:p` enfants directs
        du `w:hdr` — un tableau en en-tête (le papier à en-tête d'entreprise
        est presque toujours un tableau : logo | raison sociale | mention de
        diffusion) était perdu en entier. Le tableau doit sortir, préfixé
        `[en-tête]` comme les paragraphes, et une seule fois."""
        from docx import Document
        from docx.shared import Inches

        f = tmp_path / "header_table.docx"
        doc = Document()
        doc.add_paragraph("Corps banal sans donnee personnelle. " * 10)
        header = doc.sections[0].header
        header.is_linked_to_previous = False
        header.paragraphs[0].text = "PARA_EN_TETE"
        table = header.add_table(rows=1, cols=2, width=Inches(6))
        table.cell(0, 0).text = "Societe Exemple SA"
        table.cell(0, 1).text = "CONFIDENTIEL - DIFFUSION RESTREINTE"
        doc.save(str(f))

        result = DocxExtractor.extract(f, "header_table.docx")
        assert result.status is FileStatus.READY
        assert "[en-tête] Societe Exemple SA | CONFIDENTIEL - DIFFUSION RESTREINTE" in result.text
        assert result.text.count("CONFIDENTIEL - DIFFUSION RESTREINTE") == 1
        assert "[en-tête] PARA_EN_TETE" in result.text

    def test_empty_header_layout_row_adds_no_noise(self, tmp_path: Path) -> None:
        """D-107 : un en-tête est souvent bâti sur un tableau de mise en page
        aux cellules vides. Émettre le tableau ne doit pas produire des
        lignes `[en-tête]  | ` là où l'ancien code (paragraphes seuls)
        n'émettait rien."""
        from docx import Document
        from docx.shared import Inches

        f = tmp_path / "header_layout.docx"
        doc = Document()
        doc.add_paragraph("Corps banal sans donnee personnelle. " * 10)
        header = doc.sections[0].header
        header.is_linked_to_previous = False
        header.paragraphs[0].text = "Societe Exemple SA"
        table = header.add_table(rows=2, cols=2, width=Inches(6))
        table.cell(0, 0).text = "Logo"
        doc.save(str(f))

        text = DocxExtractor.extract(f, "header_layout.docx").text
        assert "[en-tête] Logo | " in text
        assert "[en-tête]  | " not in text

    def test_table_in_footer_is_extracted(self, tmp_path: Path) -> None:
        """D-107 : même faute côté pied de page — les gabarits RH y mettent
        la référence de traitement, le responsable et la durée de
        conservation, exactement les champs dont dépend une décision de
        suppression."""
        from docx import Document
        from docx.shared import Inches

        f = tmp_path / "footer_table.docx"
        doc = Document()
        doc.add_paragraph("Corps banal sans donnee personnelle. " * 10)
        footer = doc.sections[0].footer
        footer.is_linked_to_previous = False
        table = footer.add_table(rows=2, cols=2, width=Inches(6))
        table.cell(0, 0).text = "Reference de traitement"
        table.cell(0, 1).text = "RH-2024-017"
        table.cell(1, 0).text = "Duree de conservation : 5 ans apres depart"
        table.cell(1, 1).text = "Responsable de traitement : DRH"
        doc.save(str(f))

        result = DocxExtractor.extract(f, "footer_table.docx")
        assert result.status is FileStatus.READY
        assert "[pied de page] Reference de traitement | RH-2024-017" in result.text
        assert (
            "[pied de page] Duree de conservation : 5 ans apres depart | "
            "Responsable de traitement : DRH" in result.text
        )

    def test_comment_is_extracted(self, tmp_path: Path) -> None:
        """D-107 : `/word/comments.xml` n'était jamais lu, alors que `odf.py`
        récupère bien les `office:annotation`. Les commentaires sont le lieu
        privilégié des appréciations sur les personnes (RGPD art. 9)."""
        from docx import Document

        f = tmp_path / "comment.docx"
        doc = Document()
        para = doc.add_paragraph("Entretien annuel, appreciation globale satisfaisante. " * 5)
        doc.add_comment(
            para.runs[0],
            text="COMMENTAIRE_Mme_Dupont_arret_pour_depression_dossier_MDPH",
            author="RRH Martin",
            initials="RM",
        )
        doc.save(str(f))

        result = DocxExtractor.extract(f, "comment.docx")
        assert result.status is FileStatus.READY
        assert "[commentaires]" in result.text
        assert (
            "RRH Martin : COMMENTAIRE_Mme_Dupont_arret_pour_depression_dossier_MDPH" in result.text
        )

    def test_comment_with_id_zero_is_not_skipped(self, tmp_path: Path) -> None:
        """D-107 : `_extract_notes` ignore les `w:id` -1 et 0 (blocs
        séparateurs des notes). `comments.xml` n'a pas de tels blocs et
        numérote le premier commentaire réel `w:id="0"` — réutiliser le
        filtre des notes aurait perdu le premier commentaire de tout
        document."""
        from docx import Document

        f = tmp_path / "comment_id0.docx"
        doc = Document()
        para = doc.add_paragraph("Texte suffisamment long pour ne pas declencher LOW_TEXT. " * 5)
        doc.add_comment(para.runs[0], text="PREMIER_COMMENTAIRE_ID_ZERO", author="A")
        doc.save(str(f))

        import zipfile

        with zipfile.ZipFile(str(f)) as zf:
            assert 'w:id="0"' in zf.read("word/comments.xml").decode("utf-8")

        assert "PREMIER_COMMENTAIRE_ID_ZERO" in DocxExtractor.extract(f, "comment_id0.docx").text

    def test_sdt_wrapped_table_rows_are_extracted(self, tmp_path: Path) -> None:
        """D-107 : `Table.rows` ne rend que les `w:tr` enfants directs du
        `w:tbl`. Une ligne enveloppée dans un `w:sdt` (`CT_SdtContentRow`,
        l'encodage du contrôle « section répétitive » — le mécanisme même
        des formulaires à lignes ajoutables) disparaissait : la LLM ne
        voyait qu'un formulaire vierge (en-tête + total). LibreOffice
        restitue bien ces lignes."""
        from docx import Document

        f = tmp_path / "sdt_rows.docx"
        doc = Document()
        doc.add_paragraph("Formulaire de recensement du personnel. " * 10)
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "ENTETE_COL1"
        table.cell(0, 1).text = "ENTETE_COL2"
        table.cell(1, 0).text = "DERNIERE_LIGNE"
        table.cell(1, 1).text = "FIN"
        doc.save(str(f))

        _inject_sdt_rows(
            f,
            "DERNIERE_LIGNE",
            [
                ("LIGNE_SDT_Dupont_Jean", "LIGNE_SDT_NIR_1800175116600"),
                ("LIGNE_SDT_Martin_Sophie", "LIGNE_SDT_NIR_2790675116601"),
            ],
        )

        result = DocxExtractor.extract(f, "sdt_rows.docx")
        assert result.status is FileStatus.READY
        assert "LIGNE_SDT_Dupont_Jean | LIGNE_SDT_NIR_1800175116600" in result.text
        assert "LIGNE_SDT_Martin_Sophie | LIGNE_SDT_NIR_2790675116601" in result.text
        # Ordre du document : les lignes ajoutées restent entre l'en-tête et
        # la dernière ligne, et rien ne sort deux fois.
        assert result.text.index("ENTETE_COL1") < result.text.index("LIGNE_SDT_Dupont_Jean")
        assert result.text.index("LIGNE_SDT_Martin_Sophie") < result.text.index("DERNIERE_LIGNE")
        assert result.text.count("LIGNE_SDT_Dupont_Jean") == 1

    def test_merged_cells_output_unchanged_by_sdt_row_support(self, tmp_path: Path) -> None:
        """D-107 : le support des lignes `w:sdt` ne doit pas toucher au rendu
        des lignes ordinaires — les fusions (`gridSpan`, `vMerge`) restent
        résolues par `Table.rows` de python-docx, qui répète la cellule
        fusionnée sur chaque colonne/ligne couverte."""
        from docx import Document

        f = tmp_path / "merged.docx"
        doc = Document()
        doc.add_paragraph("Corps assez long pour ne pas declencher LOW_TEXT. " * 5)
        table = doc.add_table(rows=2, cols=2)
        for ligne in range(2):
            for col in range(2):
                table.cell(ligne, col).text = f"R{ligne}C{col}"
        table.cell(0, 0).merge(table.cell(0, 1))
        doc.save(str(f))

        text = DocxExtractor.extract(f, "merged.docx").text
        assert "R0C0\nR0C1 | R0C0\nR0C1" in text
        assert "R1C0 | R1C1" in text

    def test_footnote_is_extracted(self, tmp_path: Path) -> None:
        """Verrou de couverture : supprimer l'appel à `_extract_notes` pour
        `/word/footnotes.xml` ne faisait tomber aucun test (aucune API
        python-docx pour les notes : la partie est injectée à la main)."""
        from docx import Document

        f = tmp_path / "footnote.docx"
        doc = Document()
        doc.add_paragraph("Corps assez long pour ne pas declencher LOW_TEXT. " * 5)
        doc.save(str(f))

        _inject_part(
            f,
            "word/footnotes.xml",
            f"<w:footnotes {_W_NS}>"
            '<w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:t>SEP</w:t></w:r></w:p>'
            "</w:footnote>"
            '<w:footnote w:id="2"><w:p><w:r><w:t>TEXTE_NOTE_BAS_DE_PAGE</w:t></w:r></w:p>'
            "</w:footnote></w:footnotes>",
            f"{_OOXML_TYPE}.footnotes+xml",
            f"{_OOXML_REL}/footnotes",
        )

        result = DocxExtractor.extract(f, "footnote.docx")
        assert result.status is FileStatus.READY
        assert "[notes de bas de page]\nTEXTE_NOTE_BAS_DE_PAGE" in result.text
        assert "SEP" not in result.text  # w:id -1/0 : blocs séparateurs système

    def test_endnote_is_extracted(self, tmp_path: Path) -> None:
        """Verrou de couverture : idem pour `/word/endnotes.xml`."""
        from docx import Document

        f = tmp_path / "endnote.docx"
        doc = Document()
        doc.add_paragraph("Corps assez long pour ne pas declencher LOW_TEXT. " * 5)
        doc.save(str(f))

        _inject_part(
            f,
            "word/endnotes.xml",
            f"<w:endnotes {_W_NS}>"
            '<w:endnote w:id="3"><w:p><w:r><w:t>TEXTE_NOTE_DE_FIN</w:t></w:r></w:p>'
            "</w:endnote></w:endnotes>",
            f"{_OOXML_TYPE}.endnotes+xml",
            f"{_OOXML_REL}/endnotes",
        )

        result = DocxExtractor.extract(f, "endnote.docx")
        assert result.status is FileStatus.READY
        assert "[notes de fin]\nTEXTE_NOTE_DE_FIN" in result.text

    def test_embedded_image_export_creates_embedded_images(self, tmp_path: Path) -> None:
        """D-091 : export actif -> l'image est capturée avec un nom explicite
        et un tag `[[IMAGE: ...]]` est inséré au point d'apparition."""
        f = _docx_with_image(tmp_path, "with_image.docx", _red_square_png())

        result = DocxExtractor.extract(f, "with_image.docx", extract_images=True)
        assert result.status is FileStatus.READY
        assert len(result.embedded_images) == 1
        image = result.embedded_images[0]
        assert image.filename.startswith("with_image__img1")
        assert image.data
        assert f"[[IMAGE: {image.filename}]]" in result.text

    def test_embedded_image_export_disabled_by_default(self, tmp_path: Path) -> None:
        """D-091 : sans `extract_images`, aucune image capturée, texte inchangé
        (non-régression stricte vis-à-vis du comportement pré-D-091)."""
        f = _docx_with_image(tmp_path, "with_image2.docx", _red_square_png())

        result = DocxExtractor.extract(f, "with_image2.docx")
        assert result.embedded_images == []
        assert "[[IMAGE" not in result.text

    @pytest.mark.skipif(not _OCR_AVAILABLE, reason="Tesseract non installé")
    def test_embedded_image_ocr_extracts_text_automatically(self, tmp_path: Path) -> None:
        """D-091 : l'OCR d'une image intégrée est automatique (comme pour les
        PDF scannés), sans avoir besoin d'activer l'export."""
        from PIL import Image, ImageDraw

        img = Image.new("RGB", (600, 150), "white")
        draw = ImageDraw.Draw(img)
        draw.text((10, 50), "Bonjour le monde", fill="black")
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        f = _docx_with_image(tmp_path, "with_text_image.docx", buf.getvalue())

        result = DocxExtractor.extract(f, "with_text_image.docx")
        assert "onjour" in result.text or "monde" in result.text
        assert result.embedded_images == []
