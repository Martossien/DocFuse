"""Extracteur DOCX : .docx.

CdC §8.3 — Body, tableaux, headers/footers, footnotes, endnotes.
Utilise python-docx (MIT).
Détecte les images via word/media/* dans le ZIP.
"""

from __future__ import annotations

import logging
import zipfile
from pathlib import Path
from typing import Any

from docfuse.constants import OCR_LANG
from docfuse.core.embedded_images import build_image_marker, build_image_tag
from docfuse.core.ocr.base import OcrEngine
from docfuse.core.ocr.registry import resolve_ocr_engine
from docfuse.core.registry import register
from docfuse.extractors.base import Extractor, error_result, is_ole_encrypted, is_zip_bomb
from docfuse.i18n import t
from docfuse.models.extraction_result import EmbeddedImage, ExtractedFile
from docfuse.models.file_status import FileStatus

logger = logging.getLogger(__name__)


@register(".docx")
class DocxExtractor(Extractor):
    """Extracteur DOCX via python-docx + détection media dans le ZIP."""

    file_type = "docx"

    @classmethod
    def accepts(cls, path: Path) -> bool:
        return path.suffix.lower() == ".docx"

    @classmethod
    def extract(cls, path: Path, relative_path: str, extract_images: bool = False) -> ExtractedFile:
        try:
            # D-089 : un .docx protégé par mot de passe à l'ouverture est un
            # conteneur OLE2, plus un ZIP — sans cette détection, python-docx
            # échoue avec un `PackageNotFoundError` bas niveau qui ne dit
            # jamais à l'utilisateur que le fichier est protégé.
            if is_ole_encrypted(path):
                return ExtractedFile(
                    path=path,
                    relative_path=relative_path,
                    extension="docx",
                    file_type=cls.file_type,
                    size_bytes=path.stat().st_size,
                    status=FileStatus.ERROR,
                    error_message=t("error.encrypted_office"),
                )

            # D-093 : garde-fou "bombe zip" avant tout parsing du conteneur.
            if is_zip_bomb(path):
                return ExtractedFile(
                    path=path,
                    relative_path=relative_path,
                    extension="docx",
                    file_type=cls.file_type,
                    size_bytes=path.stat().st_size,
                    status=FileStatus.ERROR,
                    error_message=t("error.zip_bomb_suspected"),
                )

            from docx import Document
            from docx.table import Table

            # Compter les images dans word/media/
            image_count = _count_media_images(path)

            doc = Document(str(path))
            parts: list[str] = []

            # D-091 : détection des images intégrées (a:blip[@r:embed]) dans le
            # corps du document, seulement si utile (export demandé ou OCR
            # disponible) — sinon le chemin reste strictement identique à avant.
            engine = resolve_ocr_engine()
            collector = (
                _ImageCollector(doc, relative_path, extract_images, engine)
                if extract_images or engine is not None
                else None
            )

            # BUG FIX: extraire paragraphes ET tableaux dans l'ordre du document
            # python-docx n'expose pas l'ordre directement, on itère sur les enfants
            # du body element XML pour respecter l'ordre d'apparition.
            # D-069 : descend aussi dans les w:sdt (contrôles de contenu Word) —
            # sans ça, un paragraphe/tableau entier enveloppé dans un contrôle de
            # contenu (omniprésent dans les modèles RH/juridique/formulaires) est
            # invisible.
            body = doc.element.body
            parts.extend(_iter_body_parts(body, doc, Table, collector))

            # Headers et footers.
            # D-096 : un en-tête/pied « lié au précédent » (`is_linked_to_
            # previous`, le défaut de chaque nouvelle section Word) renvoie
            # la définition de la section précédente — un rapport à une
            # section par chapitre répétait le même en-tête N fois.
            for section in doc.sections:
                for header in [section.header, section.first_page_header, section.even_page_header]:
                    if header.is_linked_to_previous:
                        continue
                    for para in header.paragraphs:
                        text = _flatten_paragraph_text(para._p)
                        if text.strip():
                            parts.append(f"[en-tête] {text}")
                for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                    if footer.is_linked_to_previous:
                        continue
                    for para in footer.paragraphs:
                        text = _flatten_paragraph_text(para._p)
                        if text.strip():
                            parts.append(f"[pied de page] {text}")

            # Footnotes et endnotes — parties déjà chargées par python-docx
            # (aucune réouverture du ZIP, D-096/D-098).
            footnote_text = _extract_notes(doc, "/word/footnotes.xml", "}footnote")
            if footnote_text:
                parts.append(f"[notes de bas de page]\n{footnote_text}")

            endnote_text = _extract_notes(doc, "/word/endnotes.xml", "}endnote")
            if endnote_text:
                parts.append(f"[notes de fin]\n{endnote_text}")

            # I-19: zones de texte (w:txbxContent) du corps et des en-têtes/pieds
            textbox_text = _extract_textboxes(doc)
            if textbox_text:
                parts.append(f"[zones de texte]\n{textbox_text}")

            text = "\n\n".join(parts)

            return ExtractedFile(
                path=path,
                relative_path=relative_path,
                extension="docx",
                file_type=cls.file_type,
                size_bytes=path.stat().st_size,
                text=text,
                status=FileStatus.READY,
                image_count=image_count,
                embedded_images=collector.images if collector else [],
            )
        except Exception as exc:
            logger.exception("Erreur extraction DOCX %s", path)
            return error_result(path, relative_path, cls.file_type, exc)


def _flatten_paragraph_text(p_element: object) -> str:
    """Texte complet d'un paragraphe, en descendant dans TOUT élément imbriqué
    (`w:ins` — suivi des modifications, `w:sdt` — contrôle de contenu au
    niveau run, `w:hyperlink`, `w:fldSimple`, `w:smartTag`, ...), pas
    seulement les runs enfants directs comme le fait `Paragraph.text` de
    python-docx (D-069).

    Exclut volontairement `w:delText` (texte supprimé en suivi des
    modifications) : on veut le texte final du document, pas les deux
    versions mélangées.

    D-096 : ne descend PAS dans `w:txbxContent` (zone de texte ancrée dans
    un run du paragraphe — émise une seule fois, à part, par
    `_extract_textboxes`) ni dans `mc:Fallback` (copie VML de secours d'un
    élément DrawingML, que Word écrit en double de `mc:Choice`). Sans ces
    deux exclusions, le texte d'une zone de texte apparaissait 2 à 4 fois.
    """
    parts: list[str] = []
    stack = list(reversed(list(p_element)))  # type: ignore[call-overload]
    while stack:
        el = stack.pop()
        tag = el.tag
        if not isinstance(tag, str):  # commentaires/PI lxml
            continue
        if tag.endswith(_SKIPPED_SUBTREES):
            continue
        if tag.endswith("}t"):
            parts.append(el.text or "")
        elif tag.endswith("}tab"):
            parts.append("\t")
        elif tag.endswith(("}br", "}cr")):
            parts.append("\n")
        stack.extend(reversed(list(el)))
    return "".join(parts)


_SKIPPED_SUBTREES = ("}txbxContent", "}Fallback")


def _paragraphs_text(container: Any) -> str:
    """Texte de tous les `w:p` descendants de `container` (hors sous-arbres
    exclus), un paragraphe par ligne — jamais collés (D-096)."""
    texts: list[str] = []
    stack = list(reversed(list(container)))
    while stack:
        el = stack.pop()
        tag = el.tag
        if not isinstance(tag, str) or tag.endswith(_SKIPPED_SUBTREES):
            continue
        if tag.endswith("}p"):
            text = _flatten_paragraph_text(el)
            if text.strip():
                texts.append(text)
            continue
        stack.extend(reversed(list(el)))
    return "\n".join(texts)


class _ImageCollector:
    """Accumule les images intégrées rencontrées dans l'ordre du document (D-091).

    Résout chaque `rId` via `doc.part.rels` (relation ZIP → média), applique
    l'OCR si un moteur est disponible, et construit le marqueur inline à
    insérer au point d'apparition — jamais d'exception : une image dont la
    relation ou l'OCR échoue est simplement ignorée (pas de perte pour le
    reste du texte).
    """

    def __init__(
        self,
        doc: Any,
        relative_path: str,
        want_export: bool,
        engine: OcrEngine | None,
    ) -> None:
        self._doc = doc
        self._relative_path = relative_path
        self._want_export = want_export
        self._engine = engine
        self._count = 0
        self.images: list[EmbeddedImage] = []

    def marker_for(self, rid: str) -> str:
        rel = self._doc.part.rels.get(rid)
        if rel is None or "image" not in rel.reltype:
            return ""
        try:
            data = rel.target_part.blob
            ext = Path(rel.target_part.partname).suffix.lstrip(".") or "png"
        except Exception:
            logger.warning("Résolution de la relation image %s échouée", rid, exc_info=True)
            return ""

        self._count += 1
        tag = build_image_tag(self._relative_path, None, self._count, ext)

        ocr_text = ""
        if self._engine is not None:
            try:
                ocr_text = self._engine.ocr_image(data, OCR_LANG)
            except Exception:
                logger.warning("Échec OCR image intégrée %s", rid, exc_info=True)

        marker = build_image_marker(tag if self._want_export else None, ocr_text, OCR_LANG)
        if self._want_export and marker:
            self.images.append(EmbeddedImage(filename=tag, data=data))
        return marker


def _scan_paragraph_images(p_element: object, collector: _ImageCollector) -> str:
    """Détecte les images intégrées (`a:blip[@r:embed]`) d'un paragraphe et
    renvoie leurs marqueurs concaténés, dans l'ordre d'apparition (D-091)."""
    markers: list[str] = []
    for el in p_element.iter():  # type: ignore[attr-defined]
        if not el.tag.endswith("}blip"):
            continue
        rid = next((v for k, v in el.attrib.items() if k.endswith("}embed")), None)
        if rid:
            marker = collector.marker_for(rid)
            if marker:
                markers.append(marker)
    return "\n\n".join(markers)


def _iter_body_parts(
    container: object,
    doc: object,
    table_cls: type,
    collector: _ImageCollector | None = None,
) -> list[str]:
    """Parcourt les enfants d'un conteneur "body-like" (le corps du document,
    ou le contenu d'un `w:sdt`) et retourne paragraphes/tableaux dans l'ordre
    d'apparition — en descendant récursivement dans les `w:sdt` (contrôles de
    contenu Word) rencontrés au niveau bloc (D-069). Sans cette récursion, un
    paragraphe ou un tableau entier enveloppé dans un contrôle de contenu est
    invisible : ni `child.tag.endswith("}p")` ni `"}tbl"` ne matche `w:sdt`.

    D-083 : le contenu d'une cellule est lui-même parcouru récursivement
    (`cell._tc`, pas seulement `cell.paragraphs`) — un tableau imbriqué
    dans une cellule (fréquent dans les gabarits de rapports/formulaires
    complexes) était sinon totalement invisible : `_Cell.paragraphs` ne
    liste que les paragraphes directs, jamais un tableau imbriqué.

    D-091 : `collector`, si fourni, détecte aussi les images intégrées de
    chaque paragraphe (y compris dans les cellules de tableau) et insère
    leur marqueur juste après le texte du paragraphe qui les contient.
    """
    parts: list[str] = []
    for child in container:  # type: ignore[attr-defined]
        if child.tag.endswith("}p"):
            text = _flatten_paragraph_text(child)
            if text.strip():
                parts.append(text)
            if collector is not None:
                image_markers = _scan_paragraph_images(child, collector)
                if image_markers:
                    parts.append(image_markers)
        elif child.tag.endswith("}tbl"):
            table = table_cls(child, doc)
            for row in table.rows:
                cells = [
                    "\n".join(_iter_body_parts(cell._tc, doc, table_cls, collector)).strip()
                    for cell in row.cells
                ]
                parts.append(" | ".join(cells))
        elif child.tag.endswith("}sdt"):
            sdt_content = next((c for c in child if c.tag.endswith("}sdtContent")), None)
            if sdt_content is not None:
                parts.extend(_iter_body_parts(sdt_content, doc, table_cls, collector))
    return parts


def _count_media_images(path: Path) -> int:
    """Compte les images dans word/media/ du ZIP DOCX.

    CdC §9.4 — DOCX : word/media/* dans le ZIP.
    """
    try:
        with zipfile.ZipFile(str(path), "r") as zf:
            media_files = [n for n in zf.namelist() if n.startswith("word/media/")]
            return len(media_files)
    except Exception:
        return 0


def _part_element(doc: Any, partname: str) -> Any | None:
    """Racine lxml d'une partie du paquet DOCX déjà chargé par python-docx
    (`/word/footnotes.xml`, `/word/header1.xml`…), ou `None` si absente.

    D-096/D-098 : l'ancien code rouvrait le `.docx` (ZIP) et re-parsait la
    partie avec BeautifulSoup — jusqu'à 6 ouvertures du ZIP par fichier et
    un second parse du `document.xml` complet ~10× plus lent que celui que
    python-docx a déjà fait. Les parties XML chargées exposent `.element` ;
    les autres sont parsées une fois depuis `.blob` (lxml, jamais bs4).
    """
    from lxml import etree

    for part in doc.part.package.iter_parts():
        if str(part.partname) != partname:
            continue
        element = getattr(part, "element", None)
        if element is not None:
            return element
        try:
            return etree.fromstring(part.blob)
        except etree.XMLSyntaxError:
            logger.warning("Partie DOCX %s illisible, ignorée", partname)
            return None
    return None


def _extract_notes(doc: Any, partname: str, note_tag_suffix: str) -> str:
    """Notes de bas de page (`/word/footnotes.xml`, `}footnote`) ou de fin
    (`/word/endnotes.xml`, `}endnote`), une note par bloc, paragraphes
    séparés par `\\n` (D-096 : bs4 `get_text(strip=True)` collait runs et
    paragraphes — « …endNext » — et incluait `w:delText`, contrairement au
    corps ; `_flatten_paragraph_text` est désormais l'unique règle).

    Ignore les notes système `w:id` -1 et 0 (séparateurs).
    """
    root = _part_element(doc, partname)
    if root is None:
        return ""
    texts: list[str] = []
    for note in root:
        tag = note.tag
        if not isinstance(tag, str) or not tag.endswith(note_tag_suffix):
            continue
        note_id = next((v for k, v in note.attrib.items() if k.endswith("}id")), "")
        if note_id in ("-1", "0"):
            continue
        text = _paragraphs_text(note)
        if text:
            texts.append(text)
    return "\n".join(texts)


def _extract_textboxes(doc: Any) -> str:
    """I-19: texte des zones de texte (`w:txbxContent`) du corps et des
    en-têtes/pieds de page — l'unique endroit qui les émet (D-096).

    CdC §8.3 — DOCX : zones de texte doivent être extraites.
    D-082 : les en-têtes/pieds de page vivent dans des parties séparées
    (`word/header1.xml`…) — une zone de texte en papier à en-tête restait
    invisible tant que seul `document.xml` était lu.
    D-096 : Word 2010+ écrit chaque zone DrawingML en double — `mc:Choice`
    (wps) + `mc:Fallback` (VML) ; on ignore le sous-arbre `mc:Fallback`
    pour ne pas doubler, et les paragraphes d'une zone sont joints par
    `\\n` au lieu d'être collés.
    """
    texts: list[str] = []
    partnames = ["/word/document.xml"] + [
        str(part.partname)
        for part in doc.part.package.iter_parts()
        if str(part.partname).startswith(("/word/header", "/word/footer"))
        and str(part.partname).endswith(".xml")
    ]
    for partname in partnames:
        root = _part_element(doc, partname)
        if root is None:
            continue
        for txbx in _iter_textboxes(root):
            text = _paragraphs_text(txbx)
            if text:
                texts.append(text)
    return "\n".join(texts)


def _iter_textboxes(root: Any) -> list[Any]:
    """`w:txbxContent` d'un arbre, hors sous-arbres `mc:Fallback` (doublons
    VML) et hors zones imbriquées dans une zone déjà retenue."""
    found: list[Any] = []
    stack = [root]
    while stack:
        el = stack.pop()
        tag = el.tag
        if not isinstance(tag, str) or tag.endswith("}Fallback"):
            continue
        if tag.endswith("}txbxContent"):
            found.append(el)
            continue
        stack.extend(reversed(list(el)))
    return found
